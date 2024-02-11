import subprocess
import os
from PIL import Image, ImageFont, ImageDraw
from docx import Document
from docx.shared import Inches

PREFIX = input("Enter the prefix of the cpp files: ")
DOCX_FILENAME = input("Enter the name of the output Word document: ")
CLEANUP = input("Do you want to cleanup the text and image files after creating the Word document? (y/n): ").lower() == 'y'

# Function to save text as an image
def save_text_as_image(text, filename):
    # Set up image parameters
    image_width = 800
    margin = 20
    line_spacing = 10
    font_size = 14
    font = ImageFont.load_default()
    
    # Split text into lines
    lines = text.split('\n')
    
    # Calculate image height based on text content
    image_height = margin * 2 + (len(lines) * (font_size + line_spacing))
    
    # Create a new image
    image = Image.new('RGB', (image_width, image_height), color='white')
    draw = ImageDraw.Draw(image)
    
    # Draw text on the image
    y = margin
    for line in lines:
        draw.text((margin, y), line, fill='black', font=font)
        y += font_size + line_spacing
    
    # Save the image
    image.save(filename)

def execute_cpp(filename):
    # Compile the cpp file
    compile_process = subprocess.run(['g++', filename, '-o', 'a.exe'], capture_output=True, text=True)
    if compile_process.returncode != 0:
        print(f"Error compiling {filename}:")
        print(compile_process.stderr)
        return

    # Execute the compiled file in a separate window
    output_filename = os.path.splitext(filename)[0] + '.txt'
    execution_command = f'start cmd /k "powershell .\\a.exe > {output_filename} && exit"'
    print(execution_command)
    execute_process = subprocess.Popen(execution_command, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    execute_process.communicate()  # Wait for the program to finish execution

    # Read the text file
    with open(output_filename, 'r') as file:
        output_text = file.read()

    # Save output as an image
    image_filename = os.path.splitext(filename)[0] + '.png'
    save_text_as_image(output_text, image_filename)

    print(f"Output saved as {output_filename}")
    print(f"Output image saved as {image_filename}")

    return output_filename, image_filename

# Create a Word document
doc = Document()

# List all cpp files in the current directory
cpp_files = [file for file in os.listdir() if file.endswith('.cpp') and file.startswith(PREFIX)]

# Execute each cpp file and add code and output to the Word document
for cpp_file in cpp_files:

    #if filename.txt and filename.png already exist, continue to the next file
    if os.path.exists(os.path.splitext(cpp_file)[0] + '.txt') and os.path.exists(os.path.splitext(cpp_file)[0] + '.png'):        
         # Add code to the Word document
        doc.add_heading(f"Code: {cpp_file}", level=1)
        with open(cpp_file, 'r') as f:
            code_content = f.read()
            doc.add_paragraph(code_content)

        # Add output image to the Word document
        doc.add_heading("Output", level=2)
        doc.add_picture(os.path.splitext(cpp_file)[0] + '.png', width=Inches(6))

        # Add a page break
        doc.add_page_break()
        continue

    print(f"Executing {cpp_file}:")
    output_txt, output_img = execute_cpp(cpp_file)

    # Add code to the Word document
    doc.add_heading(f"Code: {cpp_file}", level=1)
    with open(cpp_file, 'r') as f:
        code_content = f.read()
        doc.add_paragraph(code_content)

    # Add output image to the Word document
    doc.add_heading("Output", level=2)
    doc.add_picture(output_img, width=Inches(6))

    # Add a page break
    doc.add_page_break()

    #cleanup
    if CLEANUP:
        os.remove(output_txt)
        os.remove(output_img)

# Save the Word document
doc.save(DOCX_FILENAME)
print("Word document created successfully.")
